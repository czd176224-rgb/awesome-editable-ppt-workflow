"""Extract Word pages from explicit markers or the document's physical rendered pages."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

PLUGIN_SCRIPTS = Path(__file__).resolve().parents[3] / "scripts"
if str(PLUGIN_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(PLUGIN_SCRIPTS))

from runtime_office import resolve_soffice

from docx import Document
from pypdf import PdfReader
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


DEFAULT_MARKER = r"^第\s*(\d+)\s*页(?:PPT)?$"


def _comment_catalog(document: Document) -> dict[str, dict]:
    """Return Word comments as page-local instructions, never body content."""
    catalog: dict[str, dict] = {}
    for comment in document.comments:
        text = comment.text
        if not text:
            continue
        catalog[str(comment.comment_id)] = {
            "comment_id": str(comment.comment_id),
            "text": text,
            "author": comment.author or "",
            "timestamp": comment.timestamp.isoformat() if comment.timestamp is not None else None,
        }
    return catalog


def comment_ids(block: Paragraph | Table) -> list[str]:
    """Return comment anchors/references occurring in one top-level body block."""
    values: list[str] = []
    for element in block._element.iter():
        if element.tag.rsplit("}", 1)[-1] not in {"commentRangeStart", "commentRangeEnd", "commentReference"}:
            continue
        value = element.get(qn("w:id"))
        if value is not None and value not in values:
            values.append(value)
    return values


def _bind_page_comments(pages: list[dict], catalog: dict[str, dict]) -> None:
    for page in pages:
        ordered_ids: list[str] = []
        for block in page.get("blocks", []):
            for comment_id in block.get("comment_ids", []):
                if comment_id in catalog and comment_id not in ordered_ids:
                    ordered_ids.append(comment_id)
        page["page_comments"] = [dict(catalog[comment_id]) for comment_id in ordered_ids]


def iter_blocks(document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def table_markdown(table: Table) -> str:
    def escape_cell(value: str) -> str:
        return value.strip().replace("\n", " ").replace("\\", "\\\\").replace("|", "\\|")

    rows = [[escape_cell(cell.text) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n".join(lines)


def _source_block(block: Paragraph | Table, source_block_index: int, source_order: int) -> dict:
    common = {
        "source_block_id": f"word-block-{source_block_index:06d}",
        "source_block_index": source_block_index,
        "source_order": source_order,
        "relationship_ids": relationship_ids(block),
        "comment_ids": comment_ids(block),
    }
    if isinstance(block, Table):
        return {
            **common,
            "type": "table",
            "rows": [[cell.text for cell in row.cells] for row in block.rows],
        }
    paragraph_properties = block._p.pPr
    num_properties = paragraph_properties.numPr if paragraph_properties is not None else None
    style_name = str(getattr(getattr(block, "style", None), "name", "") or "")
    if num_properties is not None or style_name.startswith("List "):
        list_kind = "number" if "Number" in style_name else "bullet"
        level = 0
        if num_properties is not None and num_properties.ilvl is not None:
            level = int(num_properties.ilvl.val)
        return {
            **common, "type": "list", "text": block.text, "list_kind": list_kind,
            "level": level, "paragraph_style": style_name,
        }
    return {**common, "type": "paragraph", "text": block.text, "paragraph_style": style_name}


def relationship_ids(block: Paragraph | Table) -> list[str]:
    """Return embedded-object relationship ids in document order for one body block."""
    relationship_ids: list[str] = []
    relationship_attributes = {
        "blip": (qn("r:embed"),),
        "imagedata": (qn("r:id"),),
        "OLEObject": (qn("r:id"),),
    }
    for element in block._element.iter():
        local_name = element.tag.rsplit("}", 1)[-1]
        for attribute in relationship_attributes.get(local_name, ()):
            relationship_id = element.get(attribute)
            if relationship_id:
                relationship_ids.append(relationship_id)
    return relationship_ids


def extract(input_path: Path, marker_pattern: str) -> dict:
    marker = re.compile(marker_pattern)
    document = Document(input_path)
    comments = _comment_catalog(document)
    pages: list[dict] = []
    current: dict | None = None
    leading_blocks: list[dict] = []
    warnings: list[dict] = []
    for source_block_index, block in enumerate(iter_blocks(document)):
        if isinstance(block, Paragraph):
            match = marker.match(block.text.strip())
            if match:
                source_page_id = int(match.group(1))
                if (
                    current is not None
                    and current["source_page_id"] == source_page_id
                    and not current["blocks"]
                ):
                    current["marker_text"] = block.text.strip()
                    current["marker_source_block_id"] = f"word-block-{source_block_index:06d}"
                    continue
                page_number = len(pages) + 1
                page_blocks = list(leading_blocks) if not pages else []
                current = {
                    "page_number": page_number,
                    "source_page_id": source_page_id,
                    "marker_text": block.text.strip(),
                    "blocks": page_blocks,
                    "must_keep": [],
                    "page_purpose": "待人工填写",
                }
                pages.append(current)
                current["marker_source_block_id"] = f"word-block-{source_block_index:06d}"
                if leading_blocks:
                    warnings.append({
                        "code": "content_before_first_marker",
                        "output_page": 1,
                        "block_count": len(leading_blocks),
                    })
                    leading_blocks.clear()
            elif current is not None:
                if block.text or relationship_ids(block) or comment_ids(block):
                    current["blocks"].append(_source_block(block, source_block_index, len(current["blocks"]) + 1))
            elif block.text or relationship_ids(block) or comment_ids(block):
                leading_blocks.append(_source_block(block, source_block_index, len(leading_blocks) + 1))
        elif isinstance(block, Table) and current is not None:
            current["blocks"].append(_source_block(block, source_block_index, len(current["blocks"]) + 1))
        elif isinstance(block, Table):
            leading_blocks.append(_source_block(block, source_block_index, len(leading_blocks) + 1))
    if not pages:
        raise ValueError("No page markers found. Expected markers such as '第1页'.")
    output_pages_by_source_id: dict[int, list[int]] = {}
    for page in pages:
        for source_order, block in enumerate(page["blocks"], start=1):
            block["source_order"] = source_order
        output_pages_by_source_id.setdefault(page["source_page_id"], []).append(page["page_number"])
        if not page["blocks"]:
            raise ValueError(f"Page {page['page_number']} is empty")
    for source_page_id in sorted(output_pages_by_source_id):
        output_pages = output_pages_by_source_id[source_page_id]
        if len(output_pages) > 1:
            warnings.append({
                "code": "duplicate_source_page_id",
                "source_page_id": source_page_id,
                "output_pages": output_pages,
            })
    _bind_page_comments(pages, comments)
    return {
        "schema_version": "1.0",
        "source_file": input_path.name,
        "page_count": len(pages),
        "pages": pages,
        "pagination_warnings": warnings,
    }


def _render_pdf_with_word(input_path: Path, output_pdf: Path) -> bool:
    try:
        import win32com.client  # type: ignore

        app = win32com.client.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        document = None
        try:
            document = app.Documents.Open(str(input_path.resolve()), ReadOnly=True, AddToRecentFiles=False)
            document.Repaginate()
            document.ExportAsFixedFormat(str(output_pdf.resolve()), 17)
            return output_pdf.is_file()
        finally:
            if document is not None:
                document.Close(False)
            app.Quit()
    except Exception:
        return False


def _render_pdf_with_libreoffice(input_path: Path, output_pdf: Path) -> bool:
    executable = resolve_soffice()
    if not executable:
        return False
    completed = subprocess.run(
        [executable, "--headless", "--convert-to", "pdf", "--outdir", str(output_pdf.parent), str(input_path)],
        capture_output=True,
        text=True,
        check=False,
        timeout=180,
    )
    generated = output_pdf.parent / f"{input_path.stem}.pdf"
    if completed.returncode == 0 and generated.is_file():
        if generated.resolve() != output_pdf.resolve():
            generated.replace(output_pdf)
        return True
    return False


def _word_block_page_evidence(input_path: Path) -> list[dict]:
    """Return Microsoft Word's page number for every top-level DOCX body block."""
    try:
        import win32com.client  # type: ignore

        app = win32com.client.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        document = None
        try:
            document = app.Documents.Open(str(input_path.resolve()), ReadOnly=True, AddToRecentFiles=False)
            document.Repaginate()
            word_blocks = []
            for paragraph in document.Paragraphs:
                block_range = paragraph.Range
                if not bool(block_range.Information(12)):  # wdWithInTable
                    word_blocks.append((int(block_range.Start), int(block_range.End), "paragraph"))
            for table in document.Tables:
                if int(table.NestingLevel) == 1:
                    block_range = table.Range
                    word_blocks.append((int(block_range.Start), int(block_range.End), "table"))
            word_blocks.sort(key=lambda value: (value[0], value[1], value[2]))
            source_blocks = list(iter_blocks(Document(input_path)))
            if len(word_blocks) != len(source_blocks):
                raise RuntimeError("Microsoft Word block/page evidence does not match DOCX body order")
            evidence = []
            for source_block_index, ((start, end, _kind), source_block) in enumerate(zip(word_blocks, source_blocks, strict=True)):
                start_range = document.Range(Start=start, End=start)
                end_position = max(start, end - 1)
                end_range = document.Range(Start=end_position, End=end_position)
                start_page = int(start_range.Information(3))  # wdActiveEndPageNumber
                end_page = int(end_range.Information(3))
                ids = relationship_ids(source_block)
                anchored_comments = comment_ids(source_block)
                evidence.append({
                    "source_block_index": source_block_index,
                    "page_number": start_page if start_page == end_page else None,
                    "relationship_ids": ids,
                    "comment_ids": anchored_comments,
                    "word_page_start": start_page,
                    "word_page_end": end_page,
                })
            return evidence
        finally:
            if document is not None:
                document.Close(False)
            app.Quit()
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError("Microsoft Word could not provide reliable block-to-page evidence") from exc


def extract_physical(input_path: Path) -> dict:
    """Render DOCX and extract text page-by-page using the physical Word pagination."""
    input_path = input_path.resolve()
    source_document = Document(input_path)
    comments = _comment_catalog(source_document)
    with tempfile.TemporaryDirectory(prefix="word-to-editable-ppt-pages-") as temporary:
        output_pdf = Path(temporary) / "rendered.pdf"
        backend = None
        if _render_pdf_with_word(input_path, output_pdf):
            backend = "microsoft-word"
        elif _render_pdf_with_libreoffice(input_path, output_pdf):
            backend = "libreoffice"
        else:
            raise RuntimeError(
                "Cannot determine physical Word pages. Install Microsoft Word or LibreOffice; marker-based guessing is not used."
            )

        block_page_evidence = _word_block_page_evidence(input_path) if backend == "microsoft-word" else None
        source_blocks = list(iter_blocks(source_document))
        if block_page_evidence is None:
            raise RuntimeError("Physical Word pagination requires exact source-block page evidence")
        spanning = [
            f"word-block-{item['source_block_index']:06d}"
            for item in block_page_evidence
            if item.get("page_number") is None
        ]
        if spanning:
            raise RuntimeError(
                "Physical Word pagination has spanning source blocks without an exact page mapping: "
                + ", ".join(spanning)
            )
        pages: list[dict] = []
        document = PdfReader(output_pdf)
        for index, page in enumerate(document.pages, start=1):
                evidence = [item for item in block_page_evidence if item.get("page_number") == index]
                blocks = [
                    _source_block(source_blocks[item["source_block_index"]], item["source_block_index"], order)
                    for order, item in enumerate(evidence, start=1)
                ]
                if not blocks:
                    raise RuntimeError(f"Physical Word page {index} has no exact source-block mapping")
                pages.append(
                    {
                        "page_number": index,
                        "blocks": blocks,
                        "must_keep": [],
                        "page_purpose": "待人工填写",
                    }
                )
    payload = {
        "schema_version": "1.1",
        "source_file": input_path.name,
        "pagination_mode": "physical_rendered_pages",
        "pagination_backend": backend,
        "page_count": len(pages),
        "pages": pages,
    }
    if block_page_evidence is not None:
        payload["block_page_evidence"] = block_page_evidence
        by_page: dict[int, list[int]] = {}
        for item in block_page_evidence:
            page_number = item.get("page_number")
            source_block_index = item.get("source_block_index")
            if isinstance(page_number, int) and isinstance(source_block_index, int):
                by_page.setdefault(page_number, []).append(source_block_index)
        for page in payload["pages"]:
            page["source_block_indexes"] = sorted(by_page.get(page["page_number"], []))
            page_comment_ids: list[str] = []
            for item in block_page_evidence:
                if item.get("page_number") != page["page_number"]:
                    continue
                for comment_id in item.get("comment_ids", []):
                    if comment_id in comments and comment_id not in page_comment_ids:
                        page_comment_ids.append(comment_id)
            page["page_comments"] = [dict(comments[comment_id]) for comment_id in page_comment_ids]
    else:
        for page in payload["pages"]:
            page["page_comments"] = []
    return payload


def extract_auto(input_path: Path, marker_pattern: str = DEFAULT_MARKER) -> dict:
    """Prefer explicit page markers; fall back to physical pages only when no markers exist."""
    try:
        payload = extract(input_path, marker_pattern)
        payload["pagination_mode"] = "explicit_text_markers"
        payload["pagination_backend"] = "docx-structure"
        return payload
    except ValueError as error:
        if not str(error).startswith("No page markers found"):
            raise
    return extract_physical(input_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--marker-regex", default=DEFAULT_MARKER)
    parser.add_argument("--pagination-mode", choices=["auto", "physical", "markers"], default="auto")
    args = parser.parse_args()
    if args.pagination_mode == "auto":
        payload = extract_auto(args.input.resolve(), args.marker_regex)
    elif args.pagination_mode == "physical":
        payload = extract_physical(args.input.resolve())
    else:
        payload = extract(args.input.resolve(), args.marker_regex)
        payload["pagination_mode"] = "explicit_text_markers"
        payload["pagination_backend"] = "docx-structure"
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"extracted_pages={payload['page_count']} output={args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
