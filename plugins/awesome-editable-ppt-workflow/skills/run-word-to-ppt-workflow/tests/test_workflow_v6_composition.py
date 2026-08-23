from __future__ import annotations

import json
import importlib.util
import sys
from copy import deepcopy
from pathlib import Path

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from workflow_v6_composition import compose_pages, validate_composition  # noqa: E402
from workflow_v6_source import initialize_v6_project  # noqa: E402


def page(source_id, *lines):
    return {
        "page_number": 1,
        "source_page_id": source_id,
        "marker_text": f"PPT第{source_id:02d}页",
        "blocks": [
            {"type": "paragraph", "text": text, "source_block_id": f"b-{source_id}-{i}"}
            for i, text in enumerate(lines, start=1)
        ],
    }


def test_compose_pages_honors_explicit_roles_and_synthesizes_missing_special_pages():
    payload = {"pages": [
        page(1, "PPT页型：封面", "项目标题", "机构名称"),
        page(3, "PPT页型：目录", "PART 1｜产业目标", "回答：产业发展路径"),
        page(4, "PPT页型：正文", "PART 1｜产业目标", "正文内容"),
        page(43, "PPT页型：正文", "结论", "最终目标：形成产业资本循环"),
    ], "pagination_warnings": []}

    result = compose_pages(payload)

    assert [item["page_role"] for item in result["pages"]] == [
        "cover", "toc", "section", "content", "content", "closing"
    ]
    assert [item["output_page_number"] for item in result["pages"]] == list(range(1, 7))
    assert result["pages"][2]["role_source"] == "synthesized"
    closing = result["pages"][-1]
    assert closing["material_source_block_ids"] == ["b-43-3"]
    blocks_by_id = {
        block["source_block_id"]: block["text"]
        for source_page in payload["pages"]
        for block in source_page["blocks"]
    }
    assert blocks_by_id[closing["material_source_block_ids"][0]] == closing["fixed_page_title"]


def test_compose_pages_infers_automatic_roles_and_preserves_pagination_warnings():
    cover = page(7, "专题报告", "呈报对象：黄石市", "联合编制")
    toc = page(9, "PART 1｜产业目标", "第2章 实施路径")
    appendix = page(15, "说明", "附录：数据口径")
    for number, item in enumerate((cover, toc, appendix), start=1):
        item["page_number"] = number
    warning = {"code": "duplicate_source_page_id", "source_page_id": 7}

    result = compose_pages({"pages": [cover, toc, appendix], "pagination_warnings": [warning]})

    assert [item["page_role"] for item in result["pages"]] == ["cover", "toc", "appendix"]
    assert [item["role_source"] for item in result["pages"]] == ["automatic"] * 3
    assert result["warnings"] == [warning]


def test_compose_pages_does_not_duplicate_an_existing_preceding_section():
    payload = {"pages": [
        page(1, "PPT页型：目录", "PART 1｜产业目标", "PART 2｜实施路径"),
        page(2, "PPT页型：章节", "PART 1｜产业目标"),
        page(3, "PPT页型：正文", "过渡说明"),
        page(4, "PPT页型：正文", "PART 1｜产业目标", "正文内容"),
    ], "pagination_warnings": []}

    result = compose_pages(payload)

    assert [item["page_role"] for item in result["pages"]].count("section") == 1


def test_synthesized_sections_from_one_toc_block_have_unique_stable_identity():
    toc = page(1, "PPT页型：目录", "PART 1｜产业目标\nPART 2｜实施路径")
    payload = {"pages": [
        toc,
        page(2, "PPT页型：正文", "PART 1｜产业目标", "正文1"),
        page(3, "PPT页型：正文", "PART 2｜实施路径", "正文2"),
    ], "pagination_warnings": []}

    result = compose_pages(payload)
    sections = [item for item in result["pages"] if item["role_source"] == "synthesized"]

    assert len(sections) == 2
    assert len({item["composition_page_id"] for item in sections}) == 2
    assert all(item["composition_page_id"].startswith("synthesized-section:") for item in sections)


def test_long_toc_splits_into_stable_continuations_without_entry_loss():
    entries = tuple(f"PART {index}｜章节{index}" for index in range(1, 27))
    payload = {"pages": [page(1, "PPT页型：目录", "目录", *entries)], "pagination_warnings": []}

    result = compose_pages(payload)
    toc_pages = [item for item in result["pages"] if item["page_role"] == "toc"]

    assert len(toc_pages) == 3
    assert [len(item["material_source_block_ids"]) for item in toc_pages] == [13, 12, 2]
    assert [
        block_id
        for item in toc_pages
        for block_id in item["material_source_block_ids"]
        if block_id != "b-1-2"
    ] == [f"b-1-{index}" for index in range(3, 29)]
    assert [item.get("composition_page_id") for item in toc_pages[1:]] == [
        "toc-continuation:b-1-2:2",
        "toc-continuation:b-1-2:3",
    ]


def test_comment_page_role_override_is_applied_and_removed_from_source_material(tmp_path: Path):
    word = tmp_path / "comment-role.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第 1 页")
    title = document.add_paragraph("普通标题")
    document.add_paragraph("正文内容")
    document.add_comment(title.runs, "PPT页型：封面", author="Reviewer", initials="RV")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)

    composition = json.loads((project / "02_v6/page_composition.json").read_text(encoding="utf-8"))
    source = json.loads((project / "02_v6/paginated_word_source.json").read_text(encoding="utf-8"))
    page_source = json.loads((project / "02_v6/page_sources/page_001.json").read_text(encoding="utf-8"))
    assert composition["pages"][0]["page_role"] == "cover"
    assert composition["pages"][0]["role_source"] == "explicit"
    assert source["pages"][0]["page_comments"] == []
    assert page_source["comments"] == []
    assert "PPT页型" not in json.dumps(page_source, ensure_ascii=False)

    server_path = SCRIPTS / "confirm_ui/server.py"
    spec = importlib.util.spec_from_file_location("comment_role_confirm_server", server_path)
    assert spec and spec.loader
    server = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(server)
    client = server.create_app(project).test_client()
    recommendations = client.get("/api/recommendations").get_json()
    template = next(
        item for item in recommendations["templates"]
        if item["id"] == recommendations["recommended_template_id"]
    )
    response = client.post("/api/confirm", json={
        "submission_id": "comment-role-audit-0001", "revision": 0,
        **template["defaults"],
        "selected_director_template_id": template["id"],
        "director_taskbook": recommendations["director_taskbook"],
    })
    assert response.status_code == 200, response.get_json()
    frozen = json.loads((project / "02_v6/page_composition.json").read_text(encoding="utf-8"))
    assert frozen["pages"][0]["role_source"] == "explicit"


def test_body_page_role_control_keeps_precedence_over_comment_control():
    raw = page(1, "PPT页型：正文", "正文标题")
    raw["page_comments"] = [{"comment_id": "c1", "text": "PPT页型：尾页"}]

    result = compose_pages({"pages": [raw], "pagination_warnings": []})

    assert result["pages"][0]["page_role"] == "content"
    assert result["pages"][0]["role_source"] == "explicit"


def test_validate_composition_rejects_duplicate_or_source_owned_synthesized_identity():
    value = compose_pages({"pages": [
        page(1, "PPT页型：目录", "PART 1｜产业目标\nPART 2｜实施路径"),
        page(2, "PPT页型：正文", "PART 1｜产业目标", "正文1"),
        page(3, "PPT页型：正文", "PART 2｜实施路径", "正文2"),
    ], "pagination_warnings": []})
    sections = [item for item in value["pages"] if item["role_source"] == "synthesized"]
    sections[1]["composition_page_id"] = sections[0]["composition_page_id"]
    with pytest.raises(ValueError, match="identity"):
        validate_composition(value)

    source_backed = next(item for item in value["pages"] if item["source_page_id"] is not None)
    sections[1]["composition_page_id"] = "synthesized-section:toc:2"
    source_backed["composition_page_id"] = "synthesized-section:forged:1"
    with pytest.raises(ValueError, match="identity"):
        validate_composition(value)

    source_backed["source_page_id"] = None
    source_backed["source_page_number"] = None
    with pytest.raises(ValueError, match="synthesized provenance"):
        validate_composition(value)


def test_validate_composition_rejects_two_covers():
    value = compose_pages({"pages": [page(1, "PPT页型：封面", "报告")], "pagination_warnings": []})
    duplicate = deepcopy(value["pages"][0])
    duplicate["output_page_number"] = 2
    value["pages"].append(duplicate)
    value["page_count"] = 2

    with pytest.raises(ValueError, match="cover"):
        validate_composition(value)


def test_validate_composition_rejects_missing_section_title():
    value = compose_pages({"pages": [page(1, "PPT页型：章节", "产业目标")], "pagination_warnings": []})
    value["pages"][0]["chapter_title"] = ""

    with pytest.raises(ValueError, match="section title"):
        validate_composition(value)


def test_validate_composition_rejects_non_continuous_output_positions():
    value = compose_pages({"pages": [page(1, "PPT页型：正文", "正文")], "pagination_warnings": []})
    value["pages"][0]["output_page_number"] = 2

    with pytest.raises(ValueError, match="continuous"):
        validate_composition(value)


def test_validate_composition_rejects_untraced_synthesized_pages():
    value = compose_pages({"pages": [
        page(1, "PPT页型：目录", "PART 1｜产业目标", "PART 2｜实施路径"),
        page(2, "PPT页型：正文", "PART 1｜产业目标", "正文"),
    ], "pagination_warnings": []})
    synthesized = next(item for item in value["pages"] if item["role_source"] == "synthesized")
    synthesized["material_source_block_ids"] = []

    with pytest.raises(ValueError, match="source block"):
        validate_composition(value)


def test_initialize_v6_project_publishes_composed_continuous_state(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    for marker, lines in (
        ("PPT第01页", ("PPT页型：封面", "专题报告", "呈报对象：黄石市")),
        ("PPT第03页", ("PPT页型：目录", "PART 1｜产业目标", "PART 2｜实施路径")),
        ("PPT第09页", ("PPT页型：正文", "PART 1｜产业目标", "正文内容")),
        ("PPT第43页", ("PPT页型：正文", "最终目标：形成产业资本循环",)),
    ):
        document.add_paragraph(marker)
        for line in lines:
            document.add_paragraph(line)
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    state = initialize_v6_project(word, logo, project)
    composition = json.loads((project / "02_v6/page_composition.json").read_text(encoding="utf-8"))

    assert [item["output_page_number"] for item in composition["pages"]] == list(range(1, 7))
    assert [item["source_page_id"] for item in composition["pages"]] == [1, 3, None, 9, 43, None]
    assert [item["page_number"] for item in state["pages"]] == list(range(1, 7))
    assert (project / "02_v6/page_sources/page_003.json").is_file()
    assert (project / "02_v6/page_materials/page_006.json").is_file()
    source = json.loads((project / "02_v6/page_sources/page_003.json").read_text(encoding="utf-8"))
    assert source["word_original"] == "PART 1｜产业目标"
    assert "PPT页型" not in json.dumps(source, ensure_ascii=False)


def test_initialize_long_toc_filters_each_continuation_to_its_material_ids(tmp_path: Path):
    word = tmp_path / "long-toc.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("PPT页型：目录")
    document.add_paragraph("目录")
    for index in range(1, 27):
        document.add_paragraph(f"PART {index}｜章节{index}")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    state = initialize_v6_project(word, logo, project)
    composition = json.loads((project / "02_v6/page_composition.json").read_text(encoding="utf-8"))
    source = json.loads((project / "02_v6/paginated_word_source.json").read_text(encoding="utf-8"))

    expected = [block_id for page in composition["pages"] for block_id in page["material_source_block_ids"]]
    actual = [block["source_block_id"] for page in source["pages"] for block in page["blocks"]]
    assert len(state["pages"]) == 3
    assert actual == expected
    assert len(actual) == len(set(actual)) == 27
    assert state["pages"][1]["title"] == "目录"
