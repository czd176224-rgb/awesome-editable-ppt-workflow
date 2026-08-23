from __future__ import annotations

import json
import hashlib
import os
import subprocess
import sys
from pathlib import Path

import pytest
from jsonschema import Draft202012Validator
from pypdf import PdfWriter


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from workflow_v6_contract import canonical_sha256, new_page, new_project  # noqa: E402
from workflow_v6_state import save  # noqa: E402
from workflow_v6_source import initialize_v6_project  # noqa: E402


COMMENTS = [
    "13个项目要显示在该页PPT中，的卢深视要重点标记",
    "请将文字内容抽象为战略逻辑模型，不要按照文字顺序排版，不要生成宣传海报，而是生成咨询公司用于管理层汇报的结构化分析页面。",
    "添加新闻稿图片，并且有王巍和李耀武讲话图片",
    "这页的企业Logo都要添加，添加方式是企业logo代替企业名称",
]


def _json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _project(tmp_path: Path) -> Path:
    project = tmp_path / "project"
    project.mkdir()
    state = new_project(
        word_source={"path": "00_source/source.docx", "sha256": "1" * 64},
        logo_source={"path": "00_source/logo.svg", "sha256": "2" * 64},
        pages=[new_page(number, title=f"第{number}页标题") for number in range(1, 5)],
    )
    visual = {
        "primary_color": "#123456",
        "secondary_color": "#ABCDEF",
        "background_color": "#FFFFFF",
        "cjk_font": "Microsoft YaHei",
        "latin_font": "Arial",
        "title_size_pt": 24,
        "body_size_pt": 14,
        "caption_size_pt": 10,
        "regional_characteristics": "",
        "visual_description": "管理层汇报风格",
    }
    state["style_confirmation"] = {"status": "confirmed", "contract": visual}
    state["confirmed_ui_revision"] = 1
    state["confirmed_ui_digest"] = canonical_sha256(visual)
    state["page_materials_status"] = "pending"
    save(project, state)

    pages = []
    for number, comment in enumerate(COMMENTS, start=1):
        pages.append(
            {
                "page_number": number,
                "blocks": [
                    {"type": "paragraph", "text": f"第{number}页标题", "source_block_id": f"p{number}-title", "source_block_index": 1, "source_order": 1, "relationship_ids": [], "comment_ids": []},
                    {"type": "paragraph", "text": f"第{number}页完整正文", "source_block_id": f"p{number}-body", "source_block_index": 2, "source_order": 2, "relationship_ids": [], "comment_ids": []},
                    {
                        "type": "list",
                        "list_kind": "number" if number % 2 == 0 else "bullet",
                        "level": 0,
                        "text": f"第{number}页列表项",
                        "source_block_id": f"p{number}-list",
                        "source_block_index": 3,
                        "source_order": 3,
                        "relationship_ids": [],
                        "comment_ids": [],
                    },
                    {
                        "type": "table",
                        "rows": [["企业", "项目"], [f"公司{number}", f"项目{number}"]],
                        "source_block_id": f"p{number}-table",
                        "source_block_index": 4,
                        "source_order": 4,
                        "relationship_ids": [],
                        "comment_ids": [],
                    },
                ],
                "fixed_page_title": f"第{number}页标题",
                "fixed_page_title_source_block_id": f"p{number}-title",
                "page_comments": [
                    {"comment_id": str(number - 1), "text": comment, "author": "用户", "timestamp": None}
                ],
            }
        )
    _json(
        project / "02_v6" / "paginated_word_source.json",
        {"schema_version": "1.0", "source_file": "source.docx", "page_count": 4, "pages": pages},
    )
    _json(
        project / "02_v6" / "source_assets.json",
        {
            "schema_version": "1.0",
            "assets": [
                {
                    "asset_id": "word_asset_001",
                    "media_type": "image/png",
                    "page_numbers": [1],
                    "relative_path": "00_source/word_assets/original/photo.png",
                    "sha256": "3" * 64,
                    "byte_size": 100,
                    "original_filename": "photo.png",
                    "generation_input": {
                        "relative_path": "00_source/word_assets/original/photo.png",
                        "sha256": "3" * 64,
                        "media_type": "image/png",
                        "derivation": "original_supported",
                    },
                },
                {
                    "asset_id": "word_asset_002",
                    "media_type": "application/pdf",
                    "page_numbers": [1],
                    "relative_path": "00_source/word_assets/original/report.pdf",
                    "sha256": "4" * 64,
                    "byte_size": 200,
                    "original_filename": "report.pdf",
                },
            ],
        },
    )
    attachment = project / "01_source_assets/00_source/word_assets/original/report.pdf"
    attachment.parent.mkdir(parents=True, exist_ok=True)
    writer = PdfWriter()
    writer.add_blank_page(width=320, height=180)
    with attachment.open("wb") as stream:
        writer.write(stream)
    manifest_path = project / "02_v6/source_assets.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["assets"][1]["sha256"] = hashlib.sha256(attachment.read_bytes()).hexdigest()
    manifest["assets"][1]["byte_size"] = attachment.stat().st_size
    _json(manifest_path, manifest)
    return project


def _walk_keys(value: object) -> set[str]:
    if isinstance(value, dict):
        return set(value) | set().union(*(_walk_keys(item) for item in value.values()), set())
    if isinstance(value, list):
        return set().union(*(_walk_keys(item) for item in value), set())
    return set()


@pytest.mark.parametrize("page_number, exact_comment", list(enumerate(COMMENTS, start=1)))
def test_collect_preserves_each_real_comment_verbatim_and_in_source_order(
    tmp_path: Path, page_number: int, exact_comment: str
):
    from awesome_page_materials import collect_page_materials

    materials = collect_page_materials(_project(tmp_path), page_number)

    assert materials["original_comments"] == [
        {"comment_id": str(page_number - 1), "source_order": 1, "text": exact_comment}
    ]
    assert set(materials) == {
        "page_number",
        "fixed_page_title",
        "complete_word_content",
        "original_comments",
        "word_images",
        "attachment_inputs",
        "visual_contract",
        "body_frame",
    }
    forbidden = {
        "unsupported_comment",
        "classification",
        "search",
        "search_requests",
        "summary",
        "degradation",
        "degradations",
    }
    assert _walk_keys(materials).isdisjoint(forbidden)


def test_collect_excludes_only_fixed_title_and_preserves_complete_body_blocks(tmp_path: Path):
    from awesome_page_materials import collect_page_materials

    materials = collect_page_materials(_project(tmp_path), 1)

    assert materials["fixed_page_title"] == "第1页标题"
    assert [block["type"] for block in materials["complete_word_content"]] == ["paragraph", "list", "table"]
    assert materials["complete_word_content"][0]["source_block_id"] == "p1-body"
    assert materials["complete_word_content"][1]["list_kind"] == "bullet"
    assert materials["complete_word_content"][2]["rows"] == [["企业", "项目"], ["公司1", "项目1"]]
    assert all(block.get("text") != "第1页标题" for block in materials["complete_word_content"])
    assert materials["word_images"] == [
        {
            "asset_id": "word_asset_001",
            "source_order": 1,
            "original_filename": "photo.png",
            "media_type": "image/png",
            "path": "01_source_assets/00_source/word_assets/original/photo.png",
            "sha256": "3" * 64,
            "byte_size": 100,
        }
    ]
    attachment = tmp_path / "project/01_source_assets/00_source/word_assets/original/report.pdf"
    assert materials["attachment_inputs"] == [
        {
            "asset_id": "word_asset_002",
            "source_order": 2,
            "original_filename": "report.pdf",
            "media_type": "application/pdf",
            "path": "01_source_assets/00_source/word_assets/original/report.pdf",
            "sha256": hashlib.sha256(attachment.read_bytes()).hexdigest(),
            "byte_size": attachment.stat().st_size,
        }
    ]


def test_attachment_rendering_preserves_unsupported_original_without_sending_it_to_renderer(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    from awesome_page_materials import _render_attachment_inputs_owned

    project = _project(tmp_path)
    unsupported = project / "01_source_assets/00_source/word_assets/original/oleObject1.bin"
    unsupported.write_bytes(b"opaque OLE authority")
    manifest_path = project / "02_v6/source_assets.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["assets"].append(
        {
            "asset_id": "word_asset_003",
            "media_type": "application/vnd.openxmlformats-officedocument.oleObject",
            "page_numbers": [1],
            "relative_path": "00_source/word_assets/original/oleObject1.bin",
            "sha256": hashlib.sha256(unsupported.read_bytes()).hexdigest(),
            "byte_size": unsupported.stat().st_size,
            "original_filename": "oleObject1.bin",
            "asset_role": "unsupported",
            "processing": "unavailable",
            "blocking": False,
        }
    )
    _json(manifest_path, manifest)

    rendered_paths: list[Path] = []

    class _Receipt:
        def to_dict(self) -> dict[str, object]:
            return {"rendered": True}

    def fake_render(
        _project: Path, _page_number: int, paths: list[Path], _lease: object
    ) -> list[_Receipt]:
        rendered_paths.extend(paths)
        return [_Receipt() for _ in paths]

    import awesome_attachment_render

    monkeypatch.setattr(awesome_attachment_render, "_render_page_attachments_owned", fake_render)
    result = _render_attachment_inputs_owned(project, 1, object())

    assert [path.name for path in rendered_paths] == ["report.pdf"]
    assert [item["original_filename"] for item in result] == ["report.pdf", "oleObject1.bin"]
    assert result[0]["render_receipt"] == {"rendered": True}
    assert "render_receipt" not in result[1]


def test_collect_uses_fixed_title_identification_from_paginated_source(tmp_path: Path):
    from awesome_page_materials import collect_page_materials

    project = _project(tmp_path)
    manifest_path = project / "02_v6" / "paginated_word_source.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["pages"][0]["fixed_page_title"] = "第1页完整正文"
    manifest["pages"][0]["fixed_page_title_source_block_id"] = "p1-body"
    _json(manifest_path, manifest)

    materials = collect_page_materials(project, 1)

    assert materials["fixed_page_title"] == "第1页完整正文"
    assert materials["complete_word_content"][0]["text"] == "第1页标题"


def test_schema_accepts_collected_materials(tmp_path: Path):
    from awesome_page_materials import collect_page_materials

    schema = json.loads((ROOT / "schemas" / "awesome_page_materials_v1.schema.json").read_text(encoding="utf-8"))
    materials = collect_page_materials(_project(tmp_path), 1)
    Draft202012Validator(schema).validate(materials)
    materials["visual_contract"].pop("regional_characteristics")
    materials["visual_contract"].pop("visual_description")
    Draft202012Validator(schema).validate(materials)


def test_cli_publishes_canonical_utf8_atomically_and_truthfully_advances_state(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    import workflow_v6_cli
    from workflow_v6_state import load

    project = _project(tmp_path)
    outputs = [project / "02_v6/awesome_page_materials" / f"page_{number:03d}.json" for number in range(1, 5)]
    for number, output in enumerate(outputs, start=1):
        monkeypatch.setattr(
            sys,
            "argv",
            [
                "workflow_v6_cli.py",
                "prepare-page-materials",
                "--project",
                str(project),
                "--page",
                str(number),
                "--out",
                str(output),
            ],
        )
        assert workflow_v6_cli.main() == 0
        published = output.read_bytes()
        value = json.loads(published.decode("utf-8"))
        if number == 1:
            receipt = value["attachment_inputs"][0]["render_receipt"]
            assert [item["page_number"] for item in receipt["pages"]] == [1]
            assert receipt["contact_sheet"]["path"].endswith("/contact_sheet.png")
        assert published == (
            json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n"
        ).encode("utf-8")
        assert not list(output.parent.glob(f".{output.name}.*.tmp"))
        state = load(project)
        assert state["pages"][number - 1]["material_state"] == "available"
        assert state["page_materials_status"] == ("confirmed" if number == 4 else "pending")


def test_invalid_materials_do_not_replace_existing_output(tmp_path: Path):
    from awesome_page_materials import publish_page_materials

    project = _project(tmp_path)
    manifest_path = project / "02_v6" / "paginated_word_source.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["pages"][0]["page_comments"][0]["comment_id"] = 7
    _json(manifest_path, manifest)
    output = project / "02_v6/awesome_page_materials/page_001.json"
    output.parent.mkdir(parents=True)
    output.write_bytes(b"existing\n")

    with pytest.raises(ValueError, match="comment identity or text is invalid"):
        publish_page_materials(project, 1, output)

    assert output.read_bytes() == b"existing\n"
    assert not list(output.parent.glob(f".{output.name}.*.tmp"))


@pytest.mark.parametrize("target", ["workflow_v6.json", "00_source/source.docx", "../outside.json"])
def test_publication_rejects_protected_or_outside_paths_before_write(tmp_path: Path, target: str):
    from awesome_page_materials import publish_page_materials

    project = _project(tmp_path)
    destination = project / target
    before = destination.read_bytes() if destination.is_file() else None
    with pytest.raises(ValueError, match="canonical project page-material path"):
        publish_page_materials(project, 1, destination)
    assert (destination.read_bytes() if destination.is_file() else None) == before


def test_real_docx_marker_extraction_preserves_lists_table_spaces_and_comment_whitespace(tmp_path: Path):
    from docx import Document
    from extract_docx_pages import extract

    word = tmp_path / "real.docx"
    document = Document()
    document.add_paragraph("第 1 页 PPT")
    title = document.add_paragraph("标题")
    bullet = document.add_paragraph("  项目甲  ", style="List Bullet")
    document.add_paragraph("项目乙", style="List Number")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "  左侧  "
    table.cell(0, 1).text = "右  侧"
    document.add_comment([bullet.runs[0]], "  保留评论两端空格  ", author="用户", initials="U")
    document.save(word)

    value = extract(word, r"^第\s*(\d+)\s*页(?:\s*PPT)?$")

    blocks = value["pages"][0]["blocks"]
    assert [block["type"] for block in blocks] == ["paragraph", "list", "list", "table"]
    assert blocks[1]["text"] == "  项目甲  "
    assert blocks[1]["list_kind"] == "bullet"
    assert blocks[2]["list_kind"] == "number"
    assert blocks[3]["rows"] == [["  左侧  ", "右  侧"]]
    assert [block["source_order"] for block in blocks] == [1, 2, 3, 4]
    assert len({block["source_block_id"] for block in blocks}) == 4
    assert value["pages"][0]["page_comments"][0]["text"] == "  保留评论两端空格  "


def test_title_exclusion_uses_source_identity_not_equal_text(tmp_path: Path):
    from awesome_page_materials import collect_page_materials

    project = _project(tmp_path)
    manifest_path = project / "02_v6/paginated_word_source.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    page = manifest["pages"][0]
    page["blocks"][0]["source_block_id"] = "block-title"
    page["blocks"][1]["source_block_id"] = "block-body"
    page["blocks"][1]["text"] = page["blocks"][0]["text"]
    page["fixed_page_title_source_block_id"] = "block-title"
    _json(manifest_path, manifest)
    materials = collect_page_materials(project, 1)
    assert [block["source_block_id"] for block in materials["complete_word_content"]][:1] == ["block-body"]


def test_confirmed_state_rejects_mutated_material_receipt(tmp_path: Path):
    import hashlib

    from awesome_page_materials import publish_page_materials
    from workflow_v6_state import load

    project = _project(tmp_path)
    outputs = [project / "02_v6/awesome_page_materials" / f"page_{number:03d}.json" for number in range(1, 5)]
    for number, output in enumerate(outputs, start=1):
        publish_page_materials(project, number, output)
    state = load(project)
    assert state["pages"][0]["material_receipt"] == {
        "schema_version": "awesome-page-materials-v1",
        "page_number": 1,
        "path": "02_v6/awesome_page_materials/page_001.json",
        "digest": hashlib.sha256(outputs[0].read_bytes()).hexdigest(),
    }
    outputs[0].write_bytes(b"{}\n")
    with pytest.raises(ValueError, match="material receipt digest"):
        load(project)


def test_repeated_identical_publication_reuses_verified_receipt(tmp_path: Path):
    from awesome_page_materials import publish_page_materials

    project = _project(tmp_path)
    output = project / "02_v6/awesome_page_materials/page_001.json"
    first = publish_page_materials(project, 1, output)
    before = output.read_bytes()
    second = publish_page_materials(project, 1, output)
    assert second == first
    assert output.read_bytes() == before


def test_repeated_publication_rejects_source_divergence(tmp_path: Path):
    from awesome_page_materials import publish_page_materials

    project = _project(tmp_path)
    output = project / "02_v6/awesome_page_materials/page_001.json"
    publish_page_materials(project, 1, output)
    manifest_path = project / "02_v6/paginated_word_source.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["pages"][0]["blocks"][1]["text"] = "changed source"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(ValueError, match="requested source materials diverge"):
        publish_page_materials(project, 1, output)


def test_publication_rejects_reparse_output_directory(tmp_path: Path):
    from awesome_page_materials import publish_page_materials

    project = _project(tmp_path)
    outside = tmp_path / "outside"
    outside.mkdir()
    output_dir = project / "02_v6/awesome_page_materials"
    try:
        output_dir.symlink_to(outside, target_is_directory=True)
    except OSError:
        pytest.skip("directory symlinks are unavailable")
    with pytest.raises(ValueError, match="reparse"):
        publish_page_materials(project, 1, output_dir / "page_001.json")
    assert list(outside.iterdir()) == []


@pytest.mark.skipif(os.name != "nt", reason="Windows junction race regression")
def test_junction_swap_before_exclusive_create_leaves_outside_empty(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import awesome_page_materials

    project = _project(tmp_path)
    outside = tmp_path / "outside"
    outside.mkdir()
    output_dir = project / "02_v6" / "awesome_page_materials"
    output = output_dir / "page_001.json"
    real_open_root = awesome_page_materials._open_project_root_handle
    swapped = False

    def swap_then_open(root: Path) -> int:
        nonlocal swapped
        if not swapped and output_dir.is_dir():
            output_dir.rmdir()
            command = (
                "$ErrorActionPreference='Stop'; "
                f"New-Item -ItemType Junction -Path '{output_dir}' -Target '{outside}' | Out-Null"
            )
            completed = subprocess.run(
                ["powershell", "-NoProfile", "-Command", command],
                capture_output=True, text=True, check=False,
            )
            assert completed.returncode == 0, completed.stdout + completed.stderr
            swapped = True
        return real_open_root(root)

    monkeypatch.setattr(awesome_page_materials, "_open_project_root_handle", swap_then_open)
    try:
        with pytest.raises(ValueError, match="escapes|reparse"):
            awesome_page_materials.publish_page_materials(project, 1, output)
        assert swapped is True
        assert list(outside.iterdir()) == []
    finally:
        if output_dir.exists():
            output_dir.rmdir()


@pytest.mark.skipif(os.name != "nt", reason="Windows junction race regression")
def test_junction_swap_never_creates_outside_even_when_cleanup_is_forbidden(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    import awesome_page_materials

    project = _project(tmp_path)
    outside = tmp_path / "outside"
    outside.mkdir()
    output_dir = project / "02_v6" / "awesome_page_materials"
    output = output_dir / "page_001.json"
    real_open_root = awesome_page_materials._open_project_root_handle
    swapped = False
    cleanup_calls = 0

    def swap_then_open(root: Path) -> int:
        nonlocal swapped
        if not swapped and output_dir.is_dir():
            output_dir.rmdir()
            command = (
                "$ErrorActionPreference='Stop'; "
                f"New-Item -ItemType Junction -Path '{output_dir}' -Target '{outside}' | Out-Null"
            )
            completed = subprocess.run(
                ["powershell", "-NoProfile", "-Command", command],
                capture_output=True, text=True, check=False,
            )
            assert completed.returncode == 0, completed.stdout + completed.stderr
            swapped = True
        return real_open_root(root)

    def forbidden_cleanup(descriptor: int, path: Path) -> None:
        nonlocal cleanup_calls
        cleanup_calls += 1
        raise PermissionError("cleanup is forbidden")

    monkeypatch.setattr(awesome_page_materials, "_open_project_root_handle", swap_then_open)
    monkeypatch.setattr(awesome_page_materials, "_delete_open_material_handle", forbidden_cleanup)
    try:
        with pytest.raises(ValueError, match="escapes|reparse"):
            awesome_page_materials.publish_page_materials(project, 1, output)
        assert swapped is True
        assert cleanup_calls == 0
        assert list(outside.iterdir()) == []
    finally:
        leaked = outside / "page_001.json"
        if leaked.exists():
            leaked.unlink()
        if output_dir.exists():
            output_dir.rmdir()


@pytest.mark.skipif(os.name != "nt", reason="Windows directory move regression")
def test_open_material_directory_denies_rename_after_containment_validation(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    import awesome_page_materials

    project = _project(tmp_path)
    outside = tmp_path / "outside"
    outside.mkdir()
    output_dir = project / "02_v6" / "awesome_page_materials"
    output = output_dir / "page_001.json"
    moved = outside / "moved-materials"
    real_verify = awesome_page_materials._verify_handle_within
    attempted = False

    def verify_then_try_move(root_descriptor: int, descriptor: int) -> None:
        nonlocal attempted
        real_verify(root_descriptor, descriptor)
        if not attempted:
            attempted = True
            completed = subprocess.run(
                [
                    "powershell", "-NoProfile", "-Command",
                    f"$ErrorActionPreference='Stop'; Move-Item -LiteralPath '{output_dir}' -Destination '{moved}'",
                ],
                capture_output=True, text=True, check=False,
            )
            assert completed.returncode != 0, "held safe directory was unexpectedly movable"

    monkeypatch.setattr(awesome_page_materials, "_verify_handle_within", verify_then_try_move)
    awesome_page_materials.publish_page_materials(project, 1, output)
    assert attempted is True
    assert output.is_file()
    assert list(outside.iterdir()) == []


@pytest.mark.skipif(os.name == "nt", reason="POSIX unsupported-platform gate")
def test_posix_publication_fails_before_any_material_write(tmp_path: Path):
    from awesome_page_materials import publish_page_materials

    project = _project(tmp_path)
    output = project / "02_v6/awesome_page_materials/page_001.json"
    with pytest.raises(RuntimeError, match="secure page-material publication is unsupported"):
        publish_page_materials(project, 1, output)
    assert not output.exists()


def test_initialize_and_collect_real_docx_keeps_exact_blocks_and_comment(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import workflow_v6_source
    from awesome_page_materials import collect_page_materials
    from docx import Document
    from workflow_v6_state import load, save

    word = tmp_path / "real.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第 1 页 PPT")
    document.add_paragraph("标题")
    item = document.add_paragraph("列表 原文", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = " 单元格 "
    table.cell(0, 1).text = "值"
    document.add_comment([item.runs[0]], " 评论保留空格 ", author="用户", initials="U")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg"/>', encoding="utf-8")
    monkeypatch.setattr(workflow_v6_source, "V6_PAGE_MARKER", r"^第\s*(\d+)\s*页(?:\s*PPT)?$")
    workflow_v6_source.initialize_v6_project(word, logo, project)
    state = load(project)
    visual_parent = tmp_path / "visual-parent"
    visual_parent.mkdir()
    visual = _project(visual_parent)
    visual_state = load(visual)
    state["style_confirmation"] = visual_state["style_confirmation"]
    state["confirmed_ui_revision"] = visual_state["confirmed_ui_revision"]
    state["confirmed_ui_digest"] = visual_state["confirmed_ui_digest"]
    state["page_materials_status"] = "pending"
    save(project, state)
    materials = collect_page_materials(project, 1)
    assert [block["type"] for block in materials["complete_word_content"]] == ["list", "table"]
    assert materials["complete_word_content"][0]["text"] == "列表 原文"
    assert materials["complete_word_content"][1]["rows"] == [[" 单元格 ", "值"]]
    assert materials["original_comments"][0]["text"] == " 评论保留空格 "


def test_comment_title_equal_to_body_never_removes_body(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import workflow_v6_source
    from awesome_page_materials import collect_page_materials
    from docx import Document
    from workflow_v6_state import load, save

    word = tmp_path / "title.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第 1 页 PPT")
    body = document.add_paragraph("DUPLICATE")
    document.add_comment([body.runs[0]], "[title: DUPLICATE]", author="用户", initials="U")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg"/>', encoding="utf-8")
    monkeypatch.setattr(workflow_v6_source, "V6_PAGE_MARKER", r"^第\s*(\d+)\s*页(?:\s*PPT)?$")
    workflow_v6_source.initialize_v6_project(word, logo, project)
    manifest = json.loads((project / "02_v6/paginated_word_source.json").read_text(encoding="utf-8"))
    assert manifest["pages"][0]["fixed_page_title_source_block_id"] is None
    state = load(project)
    visual = {"primary_color": "#123456", "secondary_color": "#ABCDEF", "background_color": "#FFFFFF", "cjk_font": "Microsoft YaHei", "latin_font": "Arial", "title_size_pt": 24, "body_size_pt": 14, "caption_size_pt": 10, "regional_characteristics": "", "visual_description": "管理层汇报风格"}
    state["style_confirmation"] = {"status": "confirmed", "contract": visual}
    state["confirmed_ui_revision"] = 1
    state["confirmed_ui_digest"] = canonical_sha256(visual)
    state["page_materials_status"] = "pending"
    save(project, state)
    assert collect_page_materials(project, 1)["complete_word_content"][0]["text"] == "DUPLICATE"


def test_table_before_first_paragraph_stays_body_and_paragraph_becomes_title(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import workflow_v6_source
    from awesome_page_materials import collect_page_materials
    from docx import Document
    from workflow_v6_state import load, save

    word = tmp_path / "table-title.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第 1 页 PPT")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Strategy Plan"
    document.add_paragraph("BODY MUST SURVIVE")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg"/>', encoding="utf-8")
    monkeypatch.setattr(workflow_v6_source, "V6_PAGE_MARKER", r"^第\s*(\d+)\s*页(?:\s*PPT)?$")
    workflow_v6_source.initialize_v6_project(word, logo, project)
    manifest = json.loads((project / "02_v6/paginated_word_source.json").read_text(encoding="utf-8"))
    page = manifest["pages"][0]
    assert page["fixed_page_title"] == "BODY MUST SURVIVE"
    assert page["fixed_page_title_source_block_id"] == page["blocks"][1]["source_block_id"]
    state = load(project)
    visual_parent = tmp_path / "visual"
    visual_parent.mkdir()
    visual_project = _project(visual_parent)
    state["style_confirmation"] = load(visual_project)["style_confirmation"]
    state["confirmed_ui_revision"] = 1
    state["confirmed_ui_digest"] = canonical_sha256(state["style_confirmation"]["contract"])
    state["page_materials_status"] = "pending"
    save(project, state)
    materials = collect_page_materials(project, 1)
    assert [block["type"] for block in materials["complete_word_content"]] == ["table"]
    assert materials["complete_word_content"][0]["rows"] == [["Strategy Plan"]]


def test_explicit_paragraph_title_identity_removes_only_that_block_with_duplicates(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import workflow_v6_source
    from awesome_page_materials import collect_page_materials
    from docx import Document
    from workflow_v6_state import load, save

    word = tmp_path / "duplicate-title.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第 1 页 PPT")
    document.add_paragraph("Strategy Plan")
    document.add_paragraph("BODY MUST SURVIVE")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Strategy Plan"
    document.add_paragraph("Strategy Plan")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg"/>', encoding="utf-8")
    monkeypatch.setattr(workflow_v6_source, "V6_PAGE_MARKER", r"^第\s*(\d+)\s*页(?:\s*PPT)?$")
    workflow_v6_source.initialize_v6_project(word, logo, project)
    manifest = json.loads((project / "02_v6/paginated_word_source.json").read_text(encoding="utf-8"))
    page = manifest["pages"][0]
    assert page["fixed_page_title_source_block_id"] == page["blocks"][0]["source_block_id"]
    state = load(project)
    visual_parent = tmp_path / "visual"
    visual_parent.mkdir()
    visual_project = _project(visual_parent)
    state["style_confirmation"] = load(visual_project)["style_confirmation"]
    state["confirmed_ui_revision"] = 1
    state["confirmed_ui_digest"] = canonical_sha256(state["style_confirmation"]["contract"])
    state["page_materials_status"] = "pending"
    save(project, state)
    blocks = collect_page_materials(project, 1)["complete_word_content"]
    assert [block["type"] for block in blocks] == ["paragraph", "table", "paragraph"]
    assert blocks[0]["text"] == "BODY MUST SURVIVE"
    assert blocks[1]["rows"] == [["Strategy Plan"]]
    assert blocks[2]["text"] == "Strategy Plan"


def test_physical_spanning_block_fails_closed(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import extract_docx_pages
    from docx import Document

    word = tmp_path / "physical.docx"
    document = Document()
    document.add_paragraph("A block spanning pages")
    document.save(word)
    monkeypatch.setattr(extract_docx_pages, "_render_pdf_with_word", lambda *_: True)
    monkeypatch.setattr(extract_docx_pages, "_word_block_page_evidence", lambda *_: [{"source_block_index": 0, "page_number": None, "word_page_start": 1, "word_page_end": 2, "comment_ids": [], "relationship_ids": []}])
    monkeypatch.setattr(extract_docx_pages, "PdfReader", lambda *_: type("Reader", (), {"pages": [object(), object()]})())
    with pytest.raises(RuntimeError, match="spanning source blocks.*word-block-000000"):
        extract_docx_pages.extract_physical(word)


def test_same_page_concurrent_publication_converges_on_one_receipt(tmp_path: Path):
    from concurrent.futures import ThreadPoolExecutor

    from awesome_page_materials import publish_page_materials
    from workflow_v6_state import load

    project = _project(tmp_path)
    output = project / "02_v6/awesome_page_materials/page_001.json"
    with ThreadPoolExecutor(max_workers=2) as pool:
        results = list(pool.map(lambda _: publish_page_materials(project, 1, output), range(2)))
    assert results[0] == results[1]
    assert load(project)["pages"][0]["material_receipt"]["digest"] == __import__("hashlib").sha256(output.read_bytes()).hexdigest()


def test_state_save_failure_rolls_back_new_material_artifact(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    import awesome_page_materials

    project = _project(tmp_path)
    output = project / "02_v6/awesome_page_materials/page_001.json"
    monkeypatch.setattr(awesome_page_materials, "save", lambda *_: (_ for _ in ()).throw(OSError("state save failed")))
    with pytest.raises(OSError, match="state save failed"):
        awesome_page_materials.publish_page_materials(project, 1, output)
    assert not output.exists()
