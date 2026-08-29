from __future__ import annotations

import hashlib
import json
import sys
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import pytest


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from extract_docx_pages import extract_auto  # noqa: E402
from workflow_v6_source import V6_PAGE_MARKER, compile_effective_page, initialize_v6_project  # noqa: E402
import workflow_v6_source  # noqa: E402
import style_recommendations  # noqa: E402


EXPECTED_DIRECTOR_TEMPLATE_IDS = [
    "company-business-introduction",
    "investment-committee",
    "project-initiation",
    "corporate-planning",
    "investment-project-bp",
]


@pytest.mark.parametrize(
    ("source_text", "expected_id"),
    [
        ("本次提交投委会审议，重点说明估值、投资回报与退出安排。", "investment-committee"),
        ("本次立项申请将说明初步尽调范围、立项依据与后续工作。", "project-initiation"),
        ("公司未来三年规划聚焦战略目标、重点任务与实施路径。", "corporate-planning"),
    ],
)
def test_director_recommendation_uses_five_scenario_templates(source_text: str, expected_id: str):
    recommendation = style_recommendations._recommendations([
        {"source_text": source_text, "page_purpose": "", "asset_bindings": []}
    ])

    assert recommendation["recommended_template_id"] == expected_id
    assert recommendation["recommendation_reason"]
    assert recommendation["recommendation_confidence"] in {"low", "medium", "high"}
    assert [item["id"] for item in recommendation["templates"]] == EXPECTED_DIRECTOR_TEMPLATE_IDS
    assert set(recommendation["director_taskbook"]) == {
        "use_scenario",
        "presenter",
        "primary_audience",
        "audience_prior_knowledge",
        "desired_outcome",
        "emphasis",
        "deemphasis",
    }


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_explicit_markers_accept_supported_forms_and_preserve_source_ids(tmp_path):
    document = Document()
    markers = [
        "第4页",
        "第 33 页",
        "第36页 · STORY LINE",
        "第26页 PPT",
        "PPT第02页",
        "PPT第44页 | PART 5｜实施路径与合作共识",
    ]
    for index, marker in enumerate(markers, start=1):
        document.add_paragraph(marker)
        document.add_paragraph(f"正文{index}")
    source = tmp_path / "markers.docx"
    document.save(source)

    payload = extract_auto(source, marker_pattern=V6_PAGE_MARKER)

    assert payload["pagination_mode"] == "explicit_text_markers"
    assert payload["page_count"] == 6
    assert [page["page_number"] for page in payload["pages"]] == [1, 2, 3, 4, 5, 6]
    assert [page["source_page_id"] for page in payload["pages"]] == [4, 33, 36, 26, 2, 44]


def test_duplicate_source_page_ids_warn_but_do_not_block(tmp_path):
    document = Document()
    for marker in ("第4页", "PPT第4页"):
        document.add_paragraph(marker)
        document.add_paragraph("正文")
    source = tmp_path / "duplicates.docx"
    document.save(source)

    payload = extract_auto(source, marker_pattern=V6_PAGE_MARKER)

    assert payload["page_count"] == 2
    assert payload["pagination_warnings"] == [{
        "code": "duplicate_source_page_id",
        "source_page_id": 4,
        "output_pages": [1, 2],
    }]


def test_adjacent_duplicate_marker_without_content_is_one_page(tmp_path):
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("第 1 页  ·  STORY LINE")
    document.add_paragraph("正文")
    source = tmp_path / "adjacent-duplicate.docx"
    document.save(source)

    payload = extract_auto(source, marker_pattern=V6_PAGE_MARKER)

    assert payload["page_count"] == 1
    assert payload["pages"][0]["marker_text"] == "第 1 页  ·  STORY LINE"
    assert [block["text"] for block in payload["pages"][0]["blocks"]] == ["正文"]
    assert payload["pagination_warnings"] == []


def test_initialize_uses_appearance_order_when_source_ids_are_9_3_9(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    for index, source_id in enumerate((9, 3, 9), start=1):
        document.add_paragraph(f"PPT第{source_id:02d}页")
        document.add_paragraph(f"正文{index}")
    document.save(word)
    logo.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>',
        encoding="utf-8",
    )

    state = initialize_v6_project(word, logo, project)
    composition = json.loads(
        (project / "02_v6/page_composition.json").read_text(encoding="utf-8")
    )

    assert [page["output_page_number"] for page in composition["pages"]] == [1, 2, 3]
    assert [page["source_page_id"] for page in composition["pages"]] == [9, 3, 9]
    assert [page["page_number"] for page in state["pages"]] == [1, 2, 3]


def test_content_before_first_marker_is_prepended_and_warned(tmp_path):
    document = Document()
    document.add_paragraph("前置说明")
    document.add_paragraph("第9页")
    document.add_paragraph("正文")
    source = tmp_path / "leading-content.docx"
    document.save(source)

    payload = extract_auto(source, marker_pattern=V6_PAGE_MARKER)

    assert [block["text"] for block in payload["pages"][0]["blocks"]] == ["前置说明", "正文"]
    assert [block["source_order"] for block in payload["pages"][0]["blocks"]] == [1, 2]
    assert payload["pagination_warnings"] == [{
        "code": "content_before_first_marker",
        "output_page": 1,
        "block_count": 1,
    }]


def test_comments_override_word_and_unavailable_attachment_invalidates_only_reference():
    value = compile_effective_page(
        page_number=1,
        word_text="原文事实为甲。",
        comments=[{"comment_id": "7", "text": "将事实改为乙，并引用附件。"}],
        references=[{"kind": "attachment", "status": "unavailable"}],
        attachment_links=[],
    )
    assert value["comment_directives"][0]["precedence"] == "overrides_word_content"
    assert value["invalidated_requirements"] == [{
        "comment_id": "7",
        "kind": "attachment_reference",
        "reason": "attachment_unavailable",
    }]


def test_search_request_records_only_page_and_purpose():
    value = compile_effective_page(
        page_number=3,
        word_text="正文",
        comments=[{"comment_id": "1", "text": "搜索相关新闻图片作为参考。"}],
        references=[],
        attachment_links=[],
    )
    assert value["search_requests"] == [
        {"page_number": 3, "purpose": "搜索相关新闻图片作为参考。"}
    ]


def test_initialize_v6_project_uses_explicit_word_pages_without_legacy_state(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("第一页标题")
    document.add_paragraph("第一页正文")
    document.add_paragraph("第2页 PPT")
    document.add_paragraph("第二页标题")
    document.add_paragraph("第二页正文")
    document.save(word)
    logo.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"></svg>',
        encoding="utf-8",
    )

    state = initialize_v6_project(word, logo, project)

    assert state["workflow_contract_version"] == "awesome-word-ppt-workflow-v1"
    assert [page["title"] for page in state["pages"]] == ["第一页标题", "第二页标题"]
    assert not (project / "workflow_run.json").exists()
    page = json.loads(
        (project / "02_v6" / "page_sources" / "page_001.json").read_text(encoding="utf-8")
    )
    assert "第一页正文" in page["word_original"]
    assert page["fixed_page_title"] == "第一页标题"
    assert page["body_render_content"] == "第一页正文"
    materials = json.loads(
        (project / "02_v6" / "page_materials" / "page_001.json").read_text(
            encoding="utf-8"
        )
    )
    assert materials["fixed_page_title"] == "第一页标题"
    assert materials["effective_body"] == "第一页正文"
    assert materials["reference_images"] == []
    assert state["page_materials_status"] == "pre_confirmation"


def test_initialize_v6_project_compiles_comment_resolution_into_confirmed_materials(tmp_path: Path):
    """Initialization must carry concrete pre-UI inputs without leaking reviewer prose to materials."""
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    title = document.add_paragraph("Growth")
    body = document.add_paragraph("Context remains. Revenue was 20%.")
    closing = document.add_paragraph("Closing remains.")
    link = document.add_paragraph()
    _add_hyperlink(link, "https://example.test/report-b", "report-b")
    document.add_comment(
        [body.runs[0]],
        "Change the revenue fact Revenue was 20% to Revenue was 30%.",
        author="Reviewer",
        initials="RV",
    )
    document.add_comment(
        [closing.runs[0]],
        "Use attachment attachment-01 rows 2, 4 fields Revenue, Margin.",
        author="Reviewer",
        initials="RV",
    )
    document.add_comment(
        [title.runs[0]],
        "[search-evidence:growth evidence]",
        author="Reviewer",
        initials="RV",
    )
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)

    materials = json.loads(
        (project / "02_v6/page_materials/page_001.json").read_text(encoding="utf-8")
    )
    expected_id = "search-request-" + hashlib.sha256(b"growth evidence").hexdigest()[:16]
    assert materials["effective_body"] == "Context remains. Revenue was 30%.\n\nClosing remains.\n\nreport-b"
    assert {item["attachment_id"] for item in materials["attachment_extracts"] if "attachment_id" in item} == {"attachment-01"}
    request = next(item for item in materials["attachment_extracts"] if item.get("attachment_id") == "attachment-01")
    assert request["selector"] == "selected_rows"
    assert request["rows"] == [2, 4]
    assert request["fields"] == ["Revenue", "Margin"]
    assert materials["image_requirements"] == [{
        "kind": "reference_acquisition", "mode": "one_shot",
        "purpose": "source_backed_evidence", "request_id": expected_id,
        "material_id": expected_id, "search_query": "growth evidence",
    }]
    assert "Change the revenue fact" not in json.dumps(materials)
    assert "Use attachment attachment-01" not in json.dumps(materials)


def test_long_first_paragraph_after_marker_is_title_authority(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    paragraph = "聚焦港澳、内地、国际三地市场，以承建地产为主业，培育实业、金融投资业务，实现利润10%增长。"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph(paragraph)
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    state = initialize_v6_project(word, logo, project)
    source = json.loads((project / "02_v6/page_sources/page_001.json").read_text(encoding="utf-8"))
    effective = json.loads((project / "02_v6/effective_pages/page_001.json").read_text(encoding="utf-8"))

    assert state["pages"][0]["title"] == paragraph
    assert source["fixed_page_title_source_block_id"] == "word-block-000001"
    assert effective["body_render_content"] == ""
    assert effective["word_original"] == paragraph


def test_explicit_body_title_instruction_is_authority_and_not_body(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("标题：黄石产业并购服务体系建设方案")
    document.add_paragraph("正文第一段")
    document.add_paragraph("正文第二段")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)
    source = json.loads((project / "02_v6/page_sources/page_001.json").read_text(encoding="utf-8"))

    assert source["fixed_page_title"] == "黄石产业并购服务体系建设方案"
    assert source["fixed_page_title_source_block_id"] == "word-block-000001"
    assert source["body_render_content"] == "正文第一段\n\n正文第二段"


def test_first_word_paragraph_precedes_later_heading_styles(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("页面提示语")
    document.add_paragraph("黄石产业并购服务体系建设方案", style="Heading 1")
    document.add_paragraph("正文内容")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    payload = extract_auto(word, marker_pattern=V6_PAGE_MARKER)
    initialize_v6_project(word, logo, project)
    source = json.loads((project / "02_v6/page_sources/page_001.json").read_text(encoding="utf-8"))
    composition = json.loads((project / "02_v6/page_composition.json").read_text(encoding="utf-8"))

    assert payload["pages"][0]["blocks"][1]["paragraph_style"] == "Heading 1"
    assert source["fixed_page_title"] == "页面提示语"
    assert source["fixed_page_title_source_block_id"] == "word-block-000001"
    assert source["body_render_content"] == "黄石产业并购服务体系建设方案\n\n正文内容"
    assert composition["pages"][0]["fixed_page_title"] == "页面提示语"


def test_page_role_control_is_not_title_authority(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("PPT页型：封面")
    document.add_paragraph("黄石产业项目建议")
    document.add_paragraph("呈报对象：黄石市")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)
    source = json.loads((project / "02_v6/page_sources/page_001.json").read_text(encoding="utf-8"))
    composition = json.loads((project / "02_v6/page_composition.json").read_text(encoding="utf-8"))

    assert source["fixed_page_title"] == "黄石产业项目建议"
    assert source["body_render_content"] == "呈报对象：黄石市"
    assert composition["pages"][0]["fixed_page_title"] == "黄石产业项目建议"


def test_initialize_v6_project_preserves_embedded_image_integrity_and_paths(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    image_path = tmp_path / "source.bmp"
    Image.new("RGB", (8, 6), color=(20, 40, 60)).save(image_path)
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("图片页")
    document.add_paragraph("正文")
    document.add_picture(str(image_path))
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)

    assets = json.loads((project / "02_v6" / "source_assets.json").read_text(encoding="utf-8"))
    asset = next(item for item in assets["assets"] if item["media_type"] == "image/bmp")
    material = json.loads(
        (project / "02_v6" / "page_materials" / "page_001.json").read_text(encoding="utf-8")
    )
    reference = material["reference_images"][0]
    generation_input = asset["generation_input"]

    assert reference["original_path"].startswith("02_v6/reference_media/word_asset_001/original.")
    assert reference["model_input_path"].startswith("02_v6/reference_media/word_asset_001/model-input.")
    assert reference["original_path"] != reference["model_input_path"]
    assert reference["integrity"]["original_sha256"] == asset["sha256"]
    assert reference["integrity"]["model_input_sha256"] != generation_input["sha256"]
    assert reference["thumbnail_path"].endswith("thumbnail.png")
    assert len(reference["integrity"]["thumbnail_sha256"]) == 64


def test_initialize_v6_project_wires_source_chart_records_as_text_facts_only(tmp_path: Path, monkeypatch):
    """Leaving extracted charts outside page materials would make chart_to_facts dead code."""
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("Page 1")
    document.add_paragraph("Chart page")
    document.add_paragraph("Narrative")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    def extract_with_chart(_word, _pages, _output):
        return {"assets": [], "chart_records": [{
            "page_numbers": [1], "title": "Revenue trend", "unit": "USD m",
            "series": [{"series": "Revenue", "time": "2025", "value": 20, "trend": "up"}],
            "image_path": "must-not-persist.png",
        }]}

    monkeypatch.setattr(workflow_v6_source, "extract_source_assets", extract_with_chart)
    monkeypatch.setattr(workflow_v6_source, "extract_auto", lambda *_args, **_kwargs: {
        "pagination_mode": "marker", "pages": [{
            "page_number": 1,
            "blocks": [
                {"type": "paragraph", "text": "Chart page", "source_block_index": 0, "source_block_id": "word-block-000000"},
                {"type": "paragraph", "text": "Narrative", "source_block_index": 1, "source_block_id": "word-block-000001"},
            ],
            "page_comments": [],
        }],
    })
    initialize_v6_project(word, logo, project)

    materials = json.loads((project / "02_v6/page_materials/page_001.json").read_text(encoding="utf-8"))
    assert materials["chart_facts"] == [{
        "title": "Revenue trend", "unit": "USD m",
        "series": [{"series": "Revenue", "time": "2025", "value": 20, "trend": "up"}],
    }]
    assert materials["reference_images"] == []


def test_real_word_schedule_table_extracts_dates_or_falls_back(tmp_path: Path):
    logo = tmp_path / "logo.svg"
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    def build_word(path: Path, end_date: str) -> None:
        document = Document()
        document.add_paragraph("第1页")
        document.add_paragraph("标题：项目计划")
        table = document.add_table(rows=3, cols=3)
        for cell, value in zip(table.rows[0].cells, ("任务", "开始日期", "结束日期")):
            cell.text = value
        for cell, value in zip(table.rows[1].cells, ("尽调", "2026-09-01", "2026-09-10")):
            cell.text = value
        for cell, value in zip(table.rows[2].cells, ("投决", "2026-09-11", end_date)):
            cell.text = value
        document.save(path)

    complete_word = tmp_path / "complete.docx"
    build_word(complete_word, "2026-09-15")
    initialize_v6_project(complete_word, logo, tmp_path / "complete-project")
    complete = json.loads(
        (tmp_path / "complete-project/02_v6/page_materials/page_001.json").read_text(encoding="utf-8")
    )["chart_facts"][0]
    assert complete["rendering_primitive"] == "time_interval"
    assert complete["series"][0]["start_dates"] == ["2026-09-01", "2026-09-11"]
    assert complete["series"][0]["end_dates"] == ["2026-09-10", "2026-09-15"]

    incomplete_word = tmp_path / "incomplete.docx"
    build_word(incomplete_word, "")
    initialize_v6_project(incomplete_word, logo, tmp_path / "incomplete-project")
    incomplete = json.loads(
        (tmp_path / "incomplete-project/02_v6/page_materials/page_001.json").read_text(encoding="utf-8")
    )["chart_facts"][0]
    assert incomplete["disabled_primitive"] == "time_interval"
    assert incomplete["fallback"] == "native_table"
    assert "rendering_primitive" not in incomplete


def test_initialize_v6_project_uses_text_derivative_for_binary_attachment(tmp_path: Path, monkeypatch):
    """Decoding an XLSX original as text would discard available extracted evidence."""
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("Page 1")
    paragraph = document.add_paragraph("Attachment evidence")
    document.add_comment(
        [paragraph.runs[0]], "Use attachment workbook rows 2 fields Revenue.",
        author="Reviewer", initials="RV",
    )
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    def extract_with_workbook(_word, _pages, output):
        original = output / "00_source/word_assets/original/workbook.xlsx"
        derivative = output / "00_source/word_assets/derived/workbook.txt"
        original.parent.mkdir(parents=True, exist_ok=True)
        derivative.parent.mkdir(parents=True, exist_ok=True)
        original.write_bytes(b"PK binary workbook")
        derivative.write_text("header\nRevenue=20\n", encoding="utf-8")
        return {"assets": [{
            "asset_id": "workbook", "page_numbers": [1],
            "media_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "relative_path": "00_source/word_assets/original/workbook.xlsx", "sha256": "a" * 64,
            "generation_input": {
                "relative_path": "00_source/word_assets/derived/workbook.txt", "sha256": "b" * 64,
                "media_type": "text/plain", "derivation": "text_extraction",
            },
        }]}

    monkeypatch.setattr(workflow_v6_source, "extract_source_assets", extract_with_workbook)
    monkeypatch.setattr(workflow_v6_source, "extract_auto", lambda *_args, **_kwargs: {
        "pagination_mode": "marker", "pages": [{
            "page_number": 1,
            "blocks": [
                {"type": "paragraph", "text": "Attachment evidence", "source_block_index": 0, "source_block_id": "word-block-000000"},
            ],
            "page_comments": [{
                "comment_id": "rows", "text": "Use attachment workbook rows 2 fields Revenue.",
            }],
        }],
    })
    initialize_v6_project(word, logo, project)

    materials = json.loads((project / "02_v6/page_materials/page_001.json").read_text(encoding="utf-8"))
    extracted = materials["attachment_extracts"][0]
    assert extracted["status"] == "available"
    assert extracted["content"] == ["Revenue=20"]
    assert extracted["source_identity"] == {
        "original_path": "01_source_assets/00_source/word_assets/original/workbook.xlsx",
        "original_sha256": "a" * 64,
    }
    assert not any(item["code"] == "attachment_unavailable" for item in materials["degradations"])
